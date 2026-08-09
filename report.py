from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from grading import calculate_grade


def generate_report(
    metadata,
    structure_score,
    objective_score,
    strengths,
    recommendations,
    section_analysis,
    intelligent_analysis,
    section_recommendations,
    filename
):
    """
    Generate a professional lesson plan review report.
    """

    document = Document()


    # ----------------------------------
    # Title
    # ----------------------------------

    title = document.add_heading(
        "EDUASSIST AI",
        level=1
    )

    title.alignment = (
        WD_PARAGRAPH_ALIGNMENT.CENTER
    )


    subtitle = document.add_heading(
        "Professional Lesson Plan Review",
        level=2
    )

    subtitle.alignment = (
        WD_PARAGRAPH_ALIGNMENT.CENTER
    )


    document.add_paragraph()


    # ----------------------------------
    # Lesson Information
    # ----------------------------------

    document.add_heading(
        "Lesson Information",
        level=2
    )


    for key, value in metadata.items():

        document.add_paragraph(
            f"{key}: {value}"
        )


    # ----------------------------------
    # Overall Lesson Quality
    # ----------------------------------

    grade, stars, readiness = calculate_grade(
        structure_score
    )


    document.add_heading(
        "Overall Lesson Quality",
        level=2
    )


    document.add_paragraph(
        f"Overall Score: {structure_score}%"
    )


    document.add_paragraph(
        f"Grade: {grade}"
    )


    document.add_paragraph(
        f"Teaching Readiness: {readiness}"
    )


    document.add_paragraph(
        f"Rating: {stars}"
    )


    document.add_paragraph(
        f"Learning Objective Score: "
        f"{objective_score}/10"
    )


    # ----------------------------------
    # Intelligent Section Quality
    # ----------------------------------

    document.add_heading(
        "Section Quality",
        level=2
    )


    for section, information in intelligent_analysis.items():

        status = information["status"]
        score = information["score"]
        feedback = information["feedback"]


        if status == "Strong":

            symbol = "🟢"

        elif status == "Needs Improvement":

            symbol = "🟡"

        elif status == "Missing":

            symbol = "🔴"

        else:

            symbol = "🔵"


        document.add_paragraph(
            f"{symbol} {section}: "
            f"{status} ({score}%)"
        )


        document.add_paragraph(
            f"Feedback: {feedback}"
        )


    # ----------------------------------
    # Strengths
    # ----------------------------------

    document.add_heading(
        "Strengths",
        level=2
    )


    for strength in strengths:

        document.add_paragraph(
            strength,
            style="List Bullet"
        )


    # ----------------------------------
    # General Recommendations
    # ----------------------------------

    document.add_heading(
        "Recommendations",
        level=2
    )


    for recommendation in recommendations:

        document.add_paragraph(
            recommendation,
            style="List Bullet"
        )


    # ----------------------------------
    # AI Professional Comment
    # ----------------------------------

    document.add_heading(
        "AI Professional Comment",
        level=2
    )


    if structure_score >= 90:

        comment = (
            "This lesson is well structured and demonstrates "
            "clear learning intentions. The lesson is suitable "
            "for classroom delivery. Continue strengthening "
            "assessment opportunities and higher-order thinking "
            "activities."
        )

    elif structure_score >= 70:

        comment = (
            "This lesson has a solid foundation. A few areas "
            "could be strengthened to improve learner engagement "
            "and assessment."
        )

    else:

        comment = (
            "This lesson requires further development before "
            "classroom delivery. Focus on improving lesson "
            "structure, learning objectives and assessment "
            "opportunities."
        )


    document.add_paragraph(
        comment
    )


    # ----------------------------------
    # Teacher Reflection
    # ----------------------------------

    document.add_heading(
        "Teacher Reflection",
        level=2
    )


    document.add_paragraph(
        "________________________________________________________"
    )

    document.add_paragraph(
        "________________________________________________________"
    )

    document.add_paragraph(
        "________________________________________________________"
    )


    # ----------------------------------
    # Section-Specific Improvement Plan
    # ----------------------------------

    document.add_heading(
        "Section-Specific Improvement Plan",
        level=2
    )


    if section_recommendations:

        for recommendation in section_recommendations:

            document.add_paragraph(
                recommendation,
                style="List Bullet"
            )

    else:

        document.add_paragraph(
            "No section-specific improvements required."
        )


    # ----------------------------------
    # Footer
    # ----------------------------------

    document.add_paragraph()


    footer = document.add_paragraph(
        "Generated by EduAssist AI\n"
        "Version 2.4\n\n"
        "Support Every Teacher • Adapt to Every School"
    )


    footer.alignment = (
        WD_PARAGRAPH_ALIGNMENT.CENTER
    )


    # ----------------------------------
    # Save Report
    # ----------------------------------

    output = (
        f"reports/Review_{filename}"
    )


    document.save(
        output
    )


    print(
        "\n✅ Professional report generated successfully!"
    )


    print(
        f"📄 Saved to: {output}"
    )